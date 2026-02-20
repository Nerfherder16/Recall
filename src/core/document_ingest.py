"""
Document ingestion pipeline.

Parses uploaded files (PDF, DOCX, markdown, text), chunks them,
extracts memories via LLM, embeds, stores, and creates
EXTRACTED_FROM edges in Neo4j.
"""

import asyncio
import hashlib
import json
import re
from typing import Any

import structlog

from .embeddings import content_hash, get_embedding_service
from .llm import LLMError, get_llm
from .models import Document, Durability, Memory, MemorySource, MemoryType

logger = structlog.get_logger()

MAX_CHUNK_CHARS = 3000

EXTRACT_PROMPT = """You are a memory extraction system. Given a chunk of a document,
extract the most important facts, decisions, or knowledge worth remembering.

Document: {filename}
Chunk {chunk_index}/{total_chunks}:
---
{chunk_text}
---

Return a JSON array of extracted memories. Each memory should have:
- "content": the fact or knowledge (1-3 sentences, standalone)
- "memory_type": one of "semantic", "procedural", "episodic"
- "importance": 1-10 (10 = critical infrastructure fact, 1 = trivial)
- "tags": array of 1-3 keyword tags

Return ONLY the JSON array. If nothing worth extracting, return [].
"""


def compute_file_hash(content: bytes) -> str:
    """SHA-256 hash of file content."""
    return hashlib.sha256(content).hexdigest()


def chunk_plaintext(text: str) -> list[str]:
    """Split plain text by paragraphs, respecting max chunk size."""
    paragraphs = re.split(r"\n\s*\n", text.strip())
    chunks = []
    current = ""

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        if len(current) + len(para) + 2 > MAX_CHUNK_CHARS and current:
            chunks.append(current.strip())
            current = para
        else:
            current = f"{current}\n\n{para}" if current else para

    if current.strip():
        chunks.append(current.strip())

    return chunks or [text[:MAX_CHUNK_CHARS]]


def chunk_markdown(text: str) -> list[str]:
    """Split markdown by headings (# / ## / ###)."""
    sections = re.split(r"(?=^#{1,3}\s)", text.strip(), flags=re.MULTILINE)
    chunks = []

    for section in sections:
        section = section.strip()
        if not section:
            continue
        if len(section) <= MAX_CHUNK_CHARS:
            chunks.append(section)
        else:
            # Section too big — sub-chunk by paragraphs
            chunks.extend(chunk_plaintext(section))

    return chunks or [text[:MAX_CHUNK_CHARS]]


def chunk_pdf_pages(pages: list[str]) -> list[str]:
    """Chunk PDF pages, merging small pages."""
    chunks = []
    current = ""

    for page_text in pages:
        page_text = page_text.strip()
        if not page_text:
            continue
        if len(current) + len(page_text) + 2 > MAX_CHUNK_CHARS and current:
            chunks.append(current.strip())
            current = page_text
        else:
            current = f"{current}\n\n{page_text}" if current else page_text

    if current.strip():
        chunks.append(current.strip())

    return chunks


def parse_docx(file_bytes: bytes) -> str:
    """Parse DOCX into plain text using python-docx."""
    import io

    import docx

    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def parse_pdf(file_bytes: bytes) -> list[str]:
    """Parse PDF into page texts using pymupdf."""
    import pymupdf

    doc = pymupdf.open(stream=file_bytes, filetype="pdf")
    pages = []
    for page in doc:
        text = page.get_text()
        if text.strip():
            pages.append(text)
    doc.close()
    return pages


async def extract_memories_from_chunk(
    chunk: str,
    chunk_index: int,
    total_chunks: int,
    filename: str,
) -> list[dict[str, Any]]:
    """Use LLM to extract memories from a single chunk."""
    llm = await get_llm()

    prompt = EXTRACT_PROMPT.format(
        filename=filename,
        chunk_index=chunk_index + 1,
        total_chunks=total_chunks,
        chunk_text=chunk[:MAX_CHUNK_CHARS],
    )

    try:
        raw = await llm.generate(prompt, format_json=True, temperature=0.2)
        data = json.loads(raw)

        # Handle single dict or array
        if isinstance(data, dict):
            data = [data]
        if not isinstance(data, list):
            return []

        return [m for m in data if isinstance(m, dict) and m.get("content")]
    except (json.JSONDecodeError, LLMError) as e:
        logger.warning("chunk_extraction_failed", chunk=chunk_index, error=str(e))
        return []


async def ingest_document(
    file_bytes: bytes,
    filename: str,
    file_type: str,
    domain: str = "general",
    durability: Durability | None = None,
    user_id: int | None = None,
    username: str | None = None,
) -> tuple[Document, list[str]]:
    """
    Full document ingestion pipeline.

    Returns (document, list_of_child_memory_ids).
    """
    from src.storage import get_neo4j_store, get_qdrant_store
    from src.storage.neo4j_documents import Neo4jDocumentStore

    file_hash = compute_file_hash(file_bytes)

    # Parse file into chunks
    if file_type == "pdf":
        pages = parse_pdf(file_bytes)
        chunks = chunk_pdf_pages(pages)
    elif file_type == "docx":
        text = parse_docx(file_bytes)
        chunks = chunk_plaintext(text)
    elif file_type == "markdown":
        text = file_bytes.decode("utf-8", errors="replace")
        chunks = chunk_markdown(text)
    else:
        text = file_bytes.decode("utf-8", errors="replace")
        chunks = chunk_plaintext(text)

    if not chunks:
        raise ValueError("No content could be extracted from the file")

    # Create document record
    doc = Document(
        filename=filename,
        file_hash=file_hash,
        file_type=file_type,
        domain=domain,
        durability=durability,
        user_id=user_id,
        username=username,
    )

    # Extract memories from each chunk (sequential — Ollama is single-threaded)
    all_extractions: list[dict[str, Any]] = []
    sem = asyncio.Semaphore(1)

    for i, chunk in enumerate(chunks):
        async with sem:
            extracted = await extract_memories_from_chunk(
                chunk,
                i,
                len(chunks),
                filename,
            )
            all_extractions.extend(extracted)
            if i < len(chunks) - 1:
                await asyncio.sleep(1)  # Avoid saturating Ollama

    # Embed and store each extracted memory
    embedding_service = await get_embedding_service()
    qdrant = await get_qdrant_store()
    neo4j = await get_neo4j_store()
    doc_store = Neo4jDocumentStore(neo4j.driver)

    # Create document node in Neo4j
    await doc_store.ensure_constraints()
    await doc_store.create_document_node(doc)

    child_ids: list[str] = []

    for extraction in all_extractions:
        raw_content = extraction["content"]
        c_hash = content_hash(raw_content)

        # Map importance from 1-10 to 0.0-1.0
        raw_imp = extraction.get("importance", 5)
        importance = max(0.0, min(1.0, float(raw_imp) / 10.0))

        memory_type_str = extraction.get("memory_type", "semantic")
        try:
            mem_type = MemoryType(memory_type_str)
        except ValueError:
            mem_type = MemoryType.SEMANTIC

        tags = extraction.get("tags", [])
        if not isinstance(tags, list):
            tags = []

        memory = Memory(
            content=raw_content,
            content_hash=c_hash,
            memory_type=mem_type,
            source=MemorySource.SYSTEM,
            domain=domain,
            tags=tags,
            importance=importance,
            initial_importance=importance,
            durability=durability,
            user_id=user_id,
            username=username,
            metadata={"document_id": doc.id},
        )

        # Embed
        embedding = await embedding_service.embed(raw_content)

        # Store in Qdrant
        await qdrant.store(memory, embedding)

        # Store in Neo4j
        await neo4j.create_memory_node(memory)

        # Create EXTRACTED_FROM edge
        await doc_store.create_extracted_from_edge(memory.id, doc.id)

        child_ids.append(memory.id)

    # Update document memory count
    doc.memory_count = len(child_ids)
    await doc_store.update_document(doc.id, memory_count=len(child_ids))

    logger.info(
        "document_ingested",
        doc_id=doc.id,
        filename=filename,
        chunks=len(chunks),
        memories=len(child_ids),
    )

    return doc, child_ids
